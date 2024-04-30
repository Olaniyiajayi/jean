import os
import json
import uuid
import logging
from logging.handlers import RotatingFileHandler

import dspy
import streamlit as st

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import boto3

log_formatter = logging.Formatter("%(asctime)s %(levelname)s %(message)s")
file_handler = RotatingFileHandler("app.log", maxBytes=1024 * 1024 * 10, backupCount=10)
file_handler.setFormatter(log_formatter)

if "course_outcomes" not in st.session_state:
    st.session_state["course_outcomes"] = []

gpt4 = dspy.OpenAI(model="gpt-3.5-turbo", max_tokens=4000, model_type="chat")
dspy.settings.configure(lm=gpt4)


def parse_course_outline(data):
    course_data = json.loads(data)
    outlines = course_data["outlines"]
    units = outlines.rstrip("\n").split("\n\n")
    unit_list = []
    for i, unit in enumerate(units, start=1):
        title, description, *objectives = unit.splitlines(keepends=True)
        title = title.strip()
        description = description.strip()
        objectives = [objective.strip() for objective in objectives]
        unit_dict = {
            "unit_number": i,
            "title": title,
            "description": description,
            "objectives": objectives,
        }
        unit_list.append(unit_dict)

    return unit_list


def update_objectives(objectives):
    updated_objectives = []
    for obj in objectives[1:]:
        objective_number, description = obj.split(". ", 1)
        description = description.strip()
        updated_objective = {
            "objective_number": int(objective_number),
            "description": description,
        }
        updated_objectives.append(updated_objective)

    return updated_objectives


def create_course_outline(
    units, output_file, bucket_name, s3_key, heading, description, outcomes
):
    print("uploading to s3")
    doc = Document()
    doc.add_heading(f"Course name: {heading}", level=1).paragraph_format.alignment = (
        WD_PARAGRAPH_ALIGNMENT.CENTER
    )
    for paragraph in doc.paragraphs:
        if f"Course name: {heading}" in paragraph.text:
            for run in paragraph.runs:
                run.font.underline = True
    doc.add_heading("Course Description:", level=2)
    doc.add_paragraph(description)
    doc.add_heading("Course Outcomes:", level=2)
    doc.add_paragraph(outcomes)
    for index, unit in enumerate(units):
        doc.add_heading(unit["title"], level=1)
        if "description" in unit:
            doc.add_paragraph(unit["description"])
        if "objectives" in unit:
            for objective in unit["objectives"]:
                doc.add_heading(
                    f"Unit {index+1} Objective {objective['objective_number']}", level=2
                )
                doc.add_paragraph(objective["content"])
                if "terms_nd_definition" in objective:
                    doc.add_heading("Terms and Definitions", level=3)
                    doc.add_paragraph(objective["terms_nd_definition"])

        for section in [
            "disccussion_questions",
            "learning_activity",
            "assessment",
            "essay_questions",
            "project",
        ]:
            if section in unit:
                doc.add_heading(section.capitalize().replace("_", " "), level=2)
                doc.add_paragraph(unit[section])
    doc.save(output_file)
    s3 = boto3.client(
        "s3",
        aws_access_key_id=os.getenv("AWS_ACCESS_KEY_ID"),
        aws_secret_access_key=os.getenv("AWS_SECRET_ACCESS_KEY"),
    )
    s3.upload_file(output_file, bucket_name, s3_key)


class Promot2Outline(dspy.Signature):
    prompt = dspy.InputField()
    course_outline = dspy.OutputField(
        desc="""For each unit,
            create a unit NAME, DESCRIPTION, and a list of 2-3 OBJECTIVES for /
            each unit. The unit description should be 1-2 paragraphs. /
            Refer to the students as “students”. The Result must start /
              with a Bloom’s verb and be measurable and specific. /
              The tone should have an active and engaging tone.
            Divide the course info into the exact number of number_of_weeks,
            Generate for all number of weeks specified, do not shorten the /
            response
            """
    )


class UnitIntroduction(dspy.Signature):
    prompt = dspy.InputField()
    unit_description = dspy.OutputField(
        desc="""Create a 2-3 paragraph unit introduction discussing /
          what students will learn in the unit.
          Discuss why this information is important and how it relates to /
          their future career in computer science."""
    )


class ObjectiveContent(dspy.Signature):
    objective = dspy.InputField()
    course_content = dspy.OutputField(
        desc="""Provide an introduction to this topic and then detailed /
        information about this topic for the students to learn the material."""
    )


class Discussion4Objective(dspy.Signature):
    objective = dspy.InputField()
    discussion_questions = dspy.OutputField(
        desc="""Create a discussion question that relates to the unit /
        objective of <Unit xx objective xx> Whenever possible, /
        have students apply the concepts to their personal lives or careers."""
    )


class TermsndDefinitions(dspy.Signature):
    objective = dspy.InputField()
    terms_and_definition = dspy.OutputField(
        desc="""Create a list of terms and definitions related to objective"""
    )


class LearningActivity(dspy.Signature):
    objective = dspy.InputField()
    learning_activity = dspy.OutputField(
        desc="""Create 1 learning activity to /
        practice the objective"""
    )


class ObjectiveQuestions(dspy.Signature):
    objective = dspy.InputField()
    objective_questions_25 = dspy.OutputField(
        desc="""Create a 20-question objective multiple choice test based on /
        this unit’s content that relates to the objectives. /
        Provide the answer key and rubric."""
    )


class EssayQuestions(dspy.Signature):
    objective = dspy.InputField()
    essay_questions_10 = dspy.OutputField(
        desc="""Create a 10 question essay question based on the unit’s /
        content that related to the following objectives.
        """
    )


class ObjectiveProject(dspy.Signature):
    objective = dspy.InputField()
    project = dspy.OutputField(
        desc="""Create a project that students need to complete related /
        to the objectives. The project can be a paper, presentation, /
        research, designing a product/service, etc. Provide a rubric /
        and answer key."""
    )


def generate_question(lesson_plan):
    """Generates a outline from the provided context using dspy."""
    try:
        data = json.loads(lesson_plan)
        lesson_plan = data.get("lesson_plan")
        grade_level = lesson_plan.get("grade_level")
        subject = lesson_plan.get("subject")
        course_name = lesson_plan.get("course_name")
        course_description = lesson_plan.get("course_description")
        course_outcomes = lesson_plan.get("course_outcomes")
        number_of_weeks = lesson_plan.get("number_of_weeks")
        prompt = f"""
            <complete_response>
            You are an instructional designer and faculty member for a school /
            teaching {grade_level} {subject} courses.
            You are developing a "{grade_level}" level {subject} /
            course called {course_name}. The course description is /
            {course_description}.
            The outcomes of the course include {course_outcomes}
            Based on this information, you will divide this course into /
            {number_of_weeks} units. return complete response for the total division.
        """
        response = dspy.ChainOfThought(Promot2Outline)(prompt=prompt).course_outline
        result = parse_course_outline(json.dumps({"outlines": response}))
        for unit in result:
            unit_prompt = f"""
                Unit {unit['unit_number']} will cover the following topics: {unit['description']}
                The unit {unit['unit_number']} objectives are: {unit['objectives']}
            """
            unit["introduction"] = dspy.ChainOfThought(UnitIntroduction)(
                prompt=unit_prompt
            ).unit_description
            unit["objectives"] = update_objectives(unit["objectives"])
            unit["learning_activity"] = dspy.ChainOfThought(LearningActivity)(
                objective=f"""As a student, I need a hands-on practice /
                activity that directly engages me in practicing each of the /
                specified objectives: {unit['objectives']}"""
            ).learning_activity
            unit["disccussion_questions"] = dspy.ChainOfThought(Discussion4Objective)(
                objective=f"{unit['objectives']}"
            ).discussion_questions
            unit["assessment"] = dspy.ChainOfThought(ObjectiveQuestions)(
                objective=f"{unit['objectives']}"
            ).objective_questions_25
            unit["essay_questions"] = dspy.ChainOfThought(EssayQuestions)(
                objective=f"{unit['objectives']}"
            ).essay_questions_10
            unit["project"] = dspy.ChainOfThought(ObjectiveProject)(
                objective=f"{unit['objectives']}"
            ).project
            for objective in unit["objectives"]:
                objective["content"] = dspy.ChainOfThought(ObjectiveContent)(
                    objective=f"""Next you will /
                    create the course content for {unit['unit_number']} /
                    One of the objectives is to {objective['description']}.""",
                ).course_content
                objective["terms_nd_definition"] = dspy.ChainOfThought(
                    TermsndDefinitions
                )(objective=objective["description"]).terms_and_definition
        key = f"{str(uuid.uuid4())}.docx"
        create_course_outline(
            result,
            "course_outline.docx",
            "dev-vecul-media-assets",
            key,
            course_name,
            course_description,
            course_outcomes,
        )
        return f"dzrsteit2h2vm.cloudfront.net/{key}"
    except ValueError as e:
        st.write(e)
        return {"error": str(e)}, 401
    except Exception as e:
        st.write("Error generating question: %s", e)
        return {"error": str(e)}, 500


def main():
    # Text Input fields for lesson plan details
    grade_level = st.text_input("Grade Level")
    subject = st.text_input("Subject")
    course_name = st.text_input("Course Name")
    course_description = st.text_area("Course Description")

    with st.expander("Course Outcomes"):
        user_input = st.text_input("Enter course outcome")
        add_button = st.button("Add", key="add_button")
        if add_button:
            if len(user_input) > 0:
                st.session_state["course_outcomes"] += [user_input]
                st.write(st.session_state["course_outcomes"])

            else:
                st.session_state["course_outcomes"] = []
    course_outcomes = st.session_state["course_outcomes"]

    number_of_weeks = st.number_input("Number of Weeks")

    generate_button = st.button("Generate Course Outline")

    if generate_button:
        lesson_plan_data = {
            "lesson_plan": {
                "grade_level": grade_level,
                "subject": subject,
                "course_name": course_name,
                "course_description": course_description,
                "course_outcomes": course_outcomes,
                "number_of_weeks": number_of_weeks,
            }
        }
        lesson_plan_json = json.dumps(lesson_plan_data)

        with st.spinner("Generating Course Outline..."):
            response = generate_question(lesson_plan_json)

        # Display response (e.g., course outline URL or error message)
        if isinstance(response, str):
            message = "Success! "  # Adjust message as needed
            st.write(message)
            st.write(f"https://{response}")
        else:
            st.error("Error generating question:", response)


if __name__ == "__main__":
    # app.run(debug=True, host="0.0.0.0", port="5000")
    main()
